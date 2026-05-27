from __future__ import annotations

import csv
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


PROJECT_ROOT = Path(__file__).resolve().parents[1]
RESULTS_DIR = PROJECT_ROOT / "Results"
PLOTS_DIR = RESULTS_DIR / "plots"
OUTPUT_PATH = Path(
    os.environ.get(
        "REPORT_OUTPUT",
        PROJECT_ROOT / "Multimodal_Emotion_Recognition_Report.docx",
    )
)


HEADING_COLOR = RGBColor(31, 77, 120)
SUBHEADING_COLOR = RGBColor(46, 116, 181)
BODY_COLOR = RGBColor(0, 0, 0)
MUTED_COLOR = RGBColor(89, 89, 89)


def rows(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def set_run(run, size: int = 12, bold: bool = False, color=BODY_COLOR) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run(run, size=16, bold=True, color=HEADING_COLOR)


def add_subheading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run(run, size=14, bold=True, color=SUBHEADING_COLOR)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run(run, size=12)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        set_run(run, size=12)
    if not paragraph.runs:
        run = paragraph.add_run(text)
    else:
        paragraph.runs[0].text = text
        run = paragraph.runs[0]
    set_run(run, size=12)


def add_table(doc: Document, headers: list[str], data: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shading = cell._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E8EEF5")
        shading.append(shd)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run(run, size=12, bold=True, color=HEADING_COLOR)

    for row in data:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run(run, size=12)
            if index > 0 and len(str(value)) < 10:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def svg_to_png(svg_path: Path, png_path: Path) -> None:
    svg = svg_path.read_text(encoding="utf-8")
    width = int(re.search(r'width="(\d+)"', svg).group(1))
    height = int(re.search(r'height="(\d+)"', svg).group(1))
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()
    title_font = ImageFont.load_default()

    for line in re.findall(r"<line .*?/>", svg):
        attrs = dict(re.findall(r'(\w+)="([^"]+)"', line))
        draw.line(
            [
                (float(attrs["x1"]), float(attrs["y1"])),
                (float(attrs["x2"]), float(attrs["y2"])),
            ],
            fill=attrs.get("stroke", "#333333"),
            width=1,
        )

    for rect in re.findall(r"<rect .*?/>", svg):
        attrs = dict(re.findall(r'(\w+)="([^"]+)"', rect))
        if "x" not in attrs:
            continue
        x = float(attrs["x"])
        y = float(attrs["y"])
        w = float(attrs["width"])
        h = float(attrs["height"])
        draw.rectangle([x, y, x + w, y + h], fill=attrs.get("fill", "#ffffff"))

    for circle in re.findall(r"<circle .*?/>", svg):
        attrs = dict(re.findall(r'(\w+)="([^"]+)"', circle))
        x = float(attrs["cx"])
        y = float(attrs["cy"])
        r = float(attrs["r"])
        fill = attrs.get("fill", "#333333")
        draw.ellipse([x - r, y - r, x + r, y + r], fill=fill)

    text_items = re.findall(r'<text ([^>]*)>(.*?)</text>', svg)
    for attrs_text, value in text_items:
        attrs = dict(re.findall(r'(\w+)="([^"]+)"', attrs_text))
        if "transform" in attrs_text:
            continue
        x = float(attrs.get("x", 0))
        y = float(attrs.get("y", 0))
        size = int(float(attrs.get("size", attrs.get("font-size", 14))))
        draw.text((x, y - size), value, fill="#222222", font=title_font if size > 18 else font)

    png_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(png_path)


def add_figure(doc: Document, title: str, image_path: Path) -> None:
    add_subheading(doc, title)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Multimodal Emotion Recognition Report")
    set_run(title_run, size=16, bold=True, color=HEADING_COLOR)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Speech-only, Text-only, and Fusion Emotion Classification")
    set_run(subtitle_run, size=12, color=MUTED_COLOR)

    add_heading(doc, "A. Architecture Decisions")
    add_subheading(doc, "Preprocessing")
    add_body(
        doc,
        "Speech audio is loaded from the TESS dataset, converted to mono float audio, resampled to 16 kHz, silence-trimmed, and framed into 25 ms windows with a 10 ms hop. Text is extracted from the spoken-word part of the filename and cleaned into lowercase alphabetic tokens.",
    )
    add_body(
        doc,
        "This design keeps the implementation aligned with the provided dataset while preparing both modalities for feature extraction.",
    )

    add_subheading(doc, "Feature Extraction")
    add_body(
        doc,
        "Speech features include RMS energy, zero-crossing rate, spectral centroid, spectral bandwidth, spectral rolloff, spectral flatness, and MFCC-style mel cepstral coefficients. Text features use word unigrams and character 2-gram/3-gram counts for the TESS-trained baseline.",
    )

    add_subheading(doc, "Temporal and Contextual Modelling")
    add_body(
        doc,
        "The speech temporal representation summarizes frame-level features using mean, standard deviation, minimum, maximum, and mean absolute delta. The text baseline uses sparse n-gram counts. To improve practical text prediction accuracy, an external pretrained DistilRoBERTa emotion model is used for live UI text inference and reported separately from TESS-trained results.",
    )
    add_body(
        doc,
        "The pretrained text model used is j-hartmann/emotion-english-distilroberta-base. It was added because the text extracted from TESS filenames is usually a neutral word and does not contain enough emotional context for strong text-only classification.",
    )

    add_subheading(doc, "Fusion and Classifier")
    add_body(
        doc,
        "Fusion uses early concatenation of speech temporal features and text n-gram features. Speech-only and fusion classifiers use distance-weighted k-nearest-neighbors with validation-tuned k. The text baseline uses Multinomial Naive Bayes.",
    )

    add_heading(doc, "B. Experiments")
    accuracy = rows(RESULTS_DIR / "all_model_accuracy_table.csv")
    summary = {}
    for row in accuracy:
        summary.setdefault(row["model_variant"], {})[row["split"]] = float(row["accuracy"])
    add_table(
        doc,
        ["Model Variant", "Train Accuracy", "Test Accuracy", "Test Samples"],
        [
            ["Speech-only", f"{summary['speech_only']['train'] * 100:.2f}%", f"{summary['speech_only']['test'] * 100:.2f}%", "560"],
            ["Text-only TESS baseline", f"{summary['text_only']['train'] * 100:.2f}%", f"{summary['text_only']['test'] * 100:.2f}%", "560"],
            ["Multimodal fusion", f"{summary['fusion']['train'] * 100:.2f}%", f"{summary['fusion']['test'] * 100:.2f}%", "560"],
        ],
    )
    add_body(
        doc,
        "Speech-only and fusion achieve high test accuracy because the TESS dataset carries emotion mainly through vocal expression. The TESS-trained text baseline is weak because the transcript words are neutral and repeated across emotion classes. A pretrained external text-emotion model was therefore integrated to improve real text prediction when users provide emotional sentences.",
    )

    add_heading(doc, "C. Analysis")
    add_subheading(doc, "Easiest and Hardest Emotions")
    report = rows(RESULTS_DIR / "speech_test_classification_report.csv")
    add_table(
        doc,
        ["Emotion", "Precision", "Recall", "F1 Score"],
        [
            [
                row["emotion"],
                f"{float(row['precision']):.2f}",
                f"{float(row['recall']):.2f}",
                f"{float(row['f1_score']):.2f}",
            ]
            for row in report
        ],
    )
    add_body(
        doc,
        "The easiest emotions are angry and fear, both reaching an F1 score of 1.00. These emotions usually have stronger acoustic patterns. The hardest emotion is pleasant_surprise, which is often confused with happy because both can share high energy and bright vocal tone.",
    )

    add_subheading(doc, "When Fusion Helps Most")
    add_body(
        doc,
        "Fusion helps most when speech and text provide complementary emotional evidence. In TESS, the transcript text is usually a neutral word, so fusion mostly follows the speech signal. If richer emotional text is provided by a user, the external pretrained text model can contribute meaningfully to live text inference.",
    )

    add_subheading(doc, "Error Analysis")
    add_table(
        doc,
        ["Pipeline", "Word", "Actual", "Predicted", "Likely Reason"],
        [
            ["Speech", "hurl", "disgust", "pleasant_surprise", "High-energy disgust can resemble surprise."],
            ["Speech", "food", "happy", "pleasant_surprise", "Happy and surprise share bright prosody."],
            ["Speech", "good", "happy", "pleasant_surprise", "Positive high-energy delivery overlaps."],
            ["Speech", "dime", "pleasant_surprise", "happy", "Surprise and happiness are acoustically close."],
            ["Fusion", "hurl", "disgust", "neutral", "Borderline speech features shifted by local neighbors."],
        ],
    )

    add_subheading(doc, "Representation Separability Visualizations")
    png_dir = RESULTS_DIR / "plots_png"
    figures = [
        ("Temporal Modelling Block", PLOTS_DIR / "speech_temporal_representation.svg", png_dir / "speech_temporal_representation.png"),
        ("Contextual Modelling Block", PLOTS_DIR / "text_contextual_representation.svg", png_dir / "text_contextual_representation.png"),
        ("Fusion Block", PLOTS_DIR / "fusion_representation.svg", png_dir / "fusion_representation.png"),
    ]
    for _, svg_path, png_path in figures:
        svg_to_png(svg_path, png_path)
    for title_text, _, png_path in figures:
        add_figure(doc, title_text, png_path)

    add_body(
        doc,
        "The speech plot shows the clearest separability, the text plot shows weak separation, and the fusion plot largely follows the speech representation because speech carries the strongest emotion signal in TESS.",
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build()
